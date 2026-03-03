from pathlib import Path
from models import TutorialDocument, TutorialStep


def _read_docx(path: Path) -> TutorialDocument:
    import docx

    doc = docx.Document(str(path))
    title = path.stem
    steps: list[TutorialStep] = []
    current_heading = None
    current_content: list[str] = []

    def flush_section():
        nonlocal current_heading, current_content
        if current_heading is not None:
            content = "\n".join(current_content).strip()
            steps.append(TutorialStep(
                index=len(steps),
                title=current_heading,
                content=content,
            ))
        current_heading = None
        current_content = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            flush_section()
            current_heading = para.text.strip() or f"Section {len(steps) + 1}"
        else:
            if current_heading is None and para.text.strip():
                # Leading paragraphs before first heading become a preamble step
                current_heading = "Introduction"
            if current_heading is not None and para.text.strip():
                current_content.append(para.text.strip())

    flush_section()

    if not steps:
        # Fallback: treat entire document as one step
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        steps.append(TutorialStep(index=0, title=title, content=all_text))

    return TutorialDocument(title=title, steps=steps)


def _read_pdf(path: Path) -> TutorialDocument:
    import pdfplumber

    title = path.stem
    steps: list[TutorialStep] = []
    current_heading = None
    current_content: list[str] = []

    def flush_section():
        nonlocal current_heading, current_content
        if current_heading is not None:
            content = "\n".join(current_content).strip()
            steps.append(TutorialStep(
                index=len(steps),
                title=current_heading,
                content=content,
            ))
        current_heading = None
        current_content = []

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["size"])
            if not words:
                continue

            # Determine modal font size (most common = body text)
            sizes = [w.get("size", 0) for w in words if w.get("size")]
            if sizes:
                from collections import Counter
                modal_size = Counter(round(s) for s in sizes).most_common(1)[0][0]
            else:
                modal_size = 12

            line_buffer: list[tuple[float, str]] = []
            for word in words:
                y = round(word["top"], 1)
                existing = next((b for b in line_buffer if b[0] == y), None)
                if existing:
                    line_buffer[line_buffer.index(existing)] = (y, existing[1] + " " + word["text"])
                else:
                    line_buffer.append((y, word["text"]))

            for y, line_text in sorted(line_buffer, key=lambda x: x[0]):
                line_stripped = line_text.strip()
                if not line_stripped:
                    continue
                # Heuristic: line is a heading if its avg word size > modal + 1
                line_words = [w for w in words if line_stripped.startswith(w["text"])]
                avg_size = (
                    sum(w.get("size", modal_size) for w in line_words) / len(line_words)
                    if line_words else modal_size
                )
                is_heading = avg_size > modal_size + 1 and len(line_stripped) < 120
                if is_heading:
                    flush_section()
                    current_heading = line_stripped
                else:
                    if current_heading is None and line_stripped:
                        current_heading = "Introduction"
                    if current_heading is not None:
                        current_content.append(line_stripped)

    flush_section()

    if not steps:
        with pdfplumber.open(str(path)) as pdf:
            all_text = "\n".join(
                page.extract_text() or "" for page in pdf.pages
            ).strip()
        steps.append(TutorialStep(index=0, title=title, content=all_text))

    return TutorialDocument(title=title, steps=steps)


def read_document(path: Path) -> TutorialDocument:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    elif suffix == ".pdf":
        return _read_pdf(path)
    else:
        raise ValueError(f"Unsupported document type: {suffix}. Use .docx or .pdf")
